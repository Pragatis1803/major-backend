from flask import Flask, request, jsonify, make_response
from flask_cors import CORS
import pytesseract
from PIL import Image
import docx
import json
import pickle
import numpy as np
app = Flask(__name__)
CORS(app)

fit_value = "age:48,bp:80,sg:1.02,al:1,su:0,bgr:121,bu:36,sc:1.2,sod:137.5288,pot:4.627244,hemo:15.4,pcv:44,wc:7800,rc:5.2,rbc_normal:1,pc_normal:1,pcc_present:0,ba_present:0,htn_yes:1,dm_yes:1,cad_yes:0,appet_poor:0,pe_yes:0,ane_yes:0"

# @app.route('/')


@app.route('/api/kidney', methods=['POST'])
def extract_text():
    # Access the image file from the POST request
    filepath = request.json['file']
    filepath_actual = filepath.replace("\\", "\\\\")
    finalPath = 'D:\\project-latest\\OrganManagement\\OrganManagement-main\\Backend\\'+filepath_actual
    #finalPath = 'D:\\Priti_kidney1.docx'
    doc = docx.Document(finalPath)
    extracted_data = {}
    for paragraph in doc.paragraphs:
        # Split each paragraph into key and value
        parts = paragraph.text.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()
            extracted_data[key] = value

    json_data = json.dumps(extracted_data)
    values_array = [float(value) for value in extracted_data.values()]
    values_array.insert(0, 0)
    print(values_array)
    kidney_model = pickle.load(open("kidney_model.pkl", "rb"))
    features = [np.array(values_array)]
    prediction = kidney_model.predict(features)
    print(prediction[0])
    if prediction == 1:
        res = "FIT"
    else:
        res = "UNFIT"
    response = make_response(res, 200)
    return response


@app.route('/api/heart', methods=['POST'])
def extract_text_heart():
    # Access the image file from the POST request
    filepath = request.json['file']
    filepath_actual = filepath.replace("\\", "\\\\")
    finalPath = 'D:\\project-latest\\OrganManagement\\OrganManagement-main\\Backend\\'+filepath_actual
    doc = docx.Document(finalPath)
    extracted_data = {}
    for paragraph in doc.paragraphs:
        # Split each paragraph into key and value
        parts = paragraph.text.split(':')
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()
            extracted_data[key] = value

    json_data = json.dumps(extracted_data)
    values_array = [float(value) for value in extracted_data.values()]
    values_array.insert(0, 0)
    print(values_array)
    kidney_model = pickle.load(open("heart_model.pkl", "rb"))
    features = [np.array(values_array)]
    prediction = kidney_model.predict(features)
    print(prediction[0])
    if prediction == 1:
        res = "FIT"
    else:
        res = "UNFIT"
    response = make_response(res, 200)
    return response


if __name__ == "__main__":
    app.run(debug=True)


# value = fit_value
#     # Read the image file using PIL
#     image = Image.open(image_file)
#     pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
#     # Perform OCR on the image to extract text
#     extracted_text = pytesseract.image_to_string(image)

#     key_value_pairs = value.split(',')

#     data = {}
#     for pair in key_value_pairs:
#         key, value = pair.split(':')
#         data[key] = value

#     json_data = json.dumps(data)
#     values_array = [float(value) for value in data.values()]
#     values_array.insert(0, 0)
#     print(values_array)
#     kidney_model = pickle.load(open("kidney_model.pkl", "rb"))
#     features = [np.array(values_array)]
#     prediction = kidney_model.predict(features)
#     print(prediction[0])
#     return "done"
